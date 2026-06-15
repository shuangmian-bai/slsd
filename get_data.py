import requests
import urllib.parse
import re
import os
from bs4 import BeautifulSoup
from docx import Document
from concurrent.futures import ThreadPoolExecutor, as_completed
import threading
from urllib.parse import urlparse

# 导入解析器工厂
from parser_factory import ParserFactory

headers = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/142.0.0.0 Safari/537.36 Edg/142.0.0.0',
}

# 创建锁对象，用于线程安全地访问文件系统
lock = threading.Lock()

# 确保数据目录存在
import sys
if getattr(sys, 'frozen', False):
    # 打包后的可执行文件
    base_dir = os.path.dirname(sys.executable)
else:
    # 开发环境
    base_dir = os.path.dirname(os.path.abspath(__file__))

data_dir = os.path.join(base_dir, "data", "word")
if not os.path.exists(data_dir):
    os.makedirs(data_dir)

# 确保图片缓存目录存在
img_dir = os.path.join(base_dir, "data", "img")
if not os.path.exists(img_dir):
    os.makedirs(img_dir)

# 确保视频缓存目录存在
video_dir = os.path.join(base_dir, "data", "video")
if not os.path.exists(video_dir):
    os.makedirs(video_dir)


def download_video(url, filename_hint='video'):
    """
    下载视频文件到本地缓存目录

    Args:
        url: 视频直链 URL
        filename_hint: 文件名提示

    Returns:
        (success, local_path) 元组
    """
    try:
        # 从 URL 中提取文件扩展名
        url_path = url.split('?')[0]
        ext = os.path.splitext(url_path)[1] or '.mp4'
        safe_name = re.sub(r'[<>:"/\\|?*\x00-\x1F]', '_', filename_hint)
        filename = f"{safe_name}{ext}"
        local_path = os.path.join(video_dir, filename)

        # 如果已存在则跳过
        if os.path.exists(local_path):
            return True, local_path

        resp = requests.get(url, headers=headers, timeout=120, stream=True)
        if resp.status_code == 200:
            with open(local_path, 'wb') as f:
                for chunk in resp.iter_content(chunk_size=8192):
                    f.write(chunk)
            return True, local_path
        return False, None
    except Exception as e:
        print(f"视频下载失败: {e}")
        return False, None

# 将文章保存为Word文档
def save_to_word(article_data, filename, source_url=None):
    # 完整文件路径
    full_path = os.path.join(data_dir, filename)

    # 检查文件是否已存在
    if os.path.exists(full_path):
        return False, full_path

    # 创建Word文档
    doc = Document()

    # 添加标题
    doc.add_heading(article_data['title'], 0)

    # 添加日期和来源信息
    doc.add_paragraph(article_data['date'])

    # 添加文章出处链接
    if source_url:
        doc.add_paragraph(f"文章出自: {source_url}")

    # 添加正文内容
    content_div = article_data['content_div']

    # 遍历内容中的所有元素
    paragraph = None
    for element in content_div.descendants:
        if element.name == 'p':
            # 处理段落
            paragraph = doc.add_paragraph(element.get_text().strip())
        elif element.name == 'div':
            # 处理div中的内容
            div_text = element.get_text().strip()
            if div_text:
                paragraph = doc.add_paragraph(div_text)
        elif element.name == 'img':
            # 处理图片
            img_src = element.get('src') or element.get('data-src')
            if img_src:
                try:
                    # 已经是完整 URL 则直接使用
                    if img_src.startswith('http'):
                        img_url = img_src
                    elif img_src.startswith('/'):
                        # 根据来源域名拼接
                        if source_url:
                            from urllib.parse import urlparse
                            parsed = urlparse(source_url)
                            img_url = f"{parsed.scheme}://{parsed.netloc}{img_src}"
                        else:
                            img_url = 'https://www.hnslsdxy.com' + img_src
                    else:
                        img_url = img_src

                    # 下载图片
                    img_response = requests.get(img_url, headers=headers, timeout=30)
                    if img_response.status_code == 200:
                        # 保存图片到img文件夹
                        img_filename = os.path.basename(img_url.split('?')[0])
                        if not img_filename:
                            img_filename = 'image.jpg'

                        img_path = os.path.join(img_dir, img_filename)
                        counter = 1
                        original_img_path = img_path
                        while os.path.exists(img_path):
                            name, ext = os.path.splitext(original_img_path)
                            img_path = f"{name}_{counter}{ext}"
                            counter += 1

                        with open(img_path, 'wb') as f:
                            f.write(img_response.content)

                        # 添加图片到文档，保持原始尺寸
                        doc.add_picture(img_path)
                    else:
                        # 图片下载失败时添加提示文字
                        doc.add_paragraph(f"[图片加载失败: {img_url}]")
                except Exception as e:
                    # 出现异常时添加提示文字
                    doc.add_paragraph(f"[图片处理失败: {str(e)}]")
        elif element.name is None and element.strip() and paragraph:
            # 处理纯文本节点（可能是段落内的文本）
            if not paragraph.text:  # 如果段落为空，则添加文本
                paragraph.text = element.strip()

    # 处理文章配图（来自解析器返回的 images 列表）
    for img_info in article_data.get('images', []):
        img_url = img_info.get('url')
        if not img_url:
            continue
        try:
            img_response = requests.get(img_url, headers=headers, timeout=30)
            if img_response.status_code == 200:
                img_filename = os.path.basename(img_url.split('?')[0])
                if not img_filename:
                    img_filename = 'image.jpg'
                img_path = os.path.join(img_dir, img_filename)
                counter = 1
                original_img_path = img_path
                while os.path.exists(img_path):
                    name, ext = os.path.splitext(original_img_path)
                    img_path = f"{name}_{counter}{ext}"
                    counter += 1
                with open(img_path, 'wb') as f:
                    f.write(img_response.content)
                doc.add_picture(img_path)
                if img_info.get('desc'):
                    doc.add_paragraph(img_info['desc'])
        except Exception as e:
            doc.add_paragraph(f"[配图处理失败: {str(e)}]")

    # 处理视频
    for idx, video_info in enumerate(article_data.get('videos', []), 1):
        video_url = video_info.get('url')
        if not video_url:
            continue

        # 插入视频封面图
        thumbnail_url = video_info.get('thumbnail')
        if thumbnail_url:
            try:
                thumb_resp = requests.get(thumbnail_url, headers=headers, timeout=30)
                if thumb_resp.status_code == 200:
                    thumb_filename = f"video_thumb_{idx}.jpg"
                    thumb_path = os.path.join(img_dir, thumb_filename)
                    with open(thumb_path, 'wb') as f:
                        f.write(thumb_resp.content)
                    doc.add_picture(thumb_path)
            except Exception:
                pass

        # 下载视频到本地
        title_hint = re.sub(r'[<>:"/\\|?*\x00-\x1F]', '_', article_data['title'])
        video_name = f"{title_hint}_video{idx}"
        success, local_path = download_video(video_url, video_name)
        if success:
            duration = video_info.get('duration', 0)
            duration_str = f"（时长 {duration // 60}分{duration % 60}秒）" if duration else ""
            doc.add_paragraph(f"[视频已缓存: {local_path}]{duration_str}")
        else:
            doc.add_paragraph(f"[视频下载失败: {video_url}]")

    # 保存文档
    doc.save(full_path)
    return True, full_path

# 处理单个链接的文章获取和保存
def process_article(url):
    try:
        # 根据URL获取对应的解析器
        parser = ParserFactory.get_parser(url)

        # 没有解析器时，尝试通用方式抓取
        if not parser:
            domain = urlparse(url).netloc
            try:
                resp = requests.get(url, headers=headers, timeout=30)
                resp.encoding = 'utf-8'
                soup = BeautifulSoup(resp.text, 'html.parser')
                title = soup.title.string.strip() if soup.title and soup.title.string else "未知标题"
                # 尝试常见的正文容器
                content_div = (
                    soup.find('article')
                    or soup.find('div', class_=re.compile(r'content|article|body|main'))
                    or soup.find('body')
                )
                body = content_div.get_text(strip=True) if content_div else ""
                article_data = {
                    'title': title,
                    'date': '',
                    'body': body,
                    'content_div': content_div or soup,
                }
                filename = re.sub(r'[<>:"/\\|?*\x00-\x1F]', '_', title) + ".docx"
                with lock:
                    saved, full_path = save_to_word(article_data, filename, source_url=url)
                if saved:
                    return True, None, f"[外部] 文章 '{title}' 已保存为 {full_path}", title, domain
                else:
                    return False, None, f"[外部] 文件 {full_path} 已存在，跳过", title, domain
            except Exception as e:
                return False, str(e), f"[外部] 处理 {url} 时出错: {e}", "未知标题", domain

        # 使用对应的解析器解析页面
        article_data = parser.parse(url)

        # 根据标题生成文件名，清理非法字符
        filename = re.sub(r'[<>:"/\\|?*\x00-\x1F]', '_', article_data['title'])
        filename = f"{filename}.docx"

        # 保存为Word文档
        with lock:
            saved, full_path = save_to_word(article_data, filename, source_url=url)

        if saved:
            return True, None, f"文章 '{article_data['title']}' 已保存为 {full_path}", article_data['title'], None
        else:
            return False, None, f"文件 {full_path} 已存在，跳过保存", article_data['title'], None
    except Exception as e:
        return False, str(e), f"处理 {url} 时出错: {e}", "未知标题", None

def get_links(search_key):
    # 对搜索关键词进行URL编码
    encoded_key = urllib.parse.quote(search_key)
    
    params = {
        'statictype': '0',
        'key': encoded_key,
        'topic': '',
        'type': 'T',
        'btime': '',
        'etime': '',
        'page': '1',
        'webid': '2014052022542170611',
    }

    response = requests.get('https://www.hnslsdxy.com/api/web_search.ashx', params=params, headers=headers)
    data = response.json()
    # 获取总页数
    page = data['pagecount']
    links = []
    external_links = []  # 收集外部链接

    for i in range(page):
        i += 1
        params['page'] = str(i)
        response = requests.get('https://www.hnslsdxy.com/api/web_search.ashx', params=params, headers=headers)
        data = response.json()
        for item in data['content']:
            link  = item['title'].split('href=')[2]
            link  = link.split("'")[1]
            if link[0:4] == 'http':
                # 检查该链接是否属于我们已注册解析器支持的域名
                if ParserFactory.is_registered_domain(link):
                    # 如果是已注册的域名，则将其作为内部链接处理
                    links.append(link)
                else:
                    # 否则将其作为真正的外部链接收集
                    external_links.append(link)
                continue
            link = 'https://www.hnslsdxy.com/page/' +  link
            links.append(link)
    
    return links, external_links